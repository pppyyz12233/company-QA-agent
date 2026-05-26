import docx
import pdfplumber

from utils.config import md5_path
from utils.handle_text import kb_service


def pdf_service(file_path: str, filename: str,operator: str = "admin"):
    print(f"开始解析 Word: {filename}, 上传者: {operator}")
    full_text=[]

    with pdfplumber.open(file_path) as pdf:
        for page_num,page in enumerate(pdf.pages, start=1):
            text = page.extract_text()#每页的内容

            if text:
                full_text.append(text)
                print(f"已经读取第{page_num}页")
            else:print(f"pdf第{page_num}页没有文字")

    if not full_text:
        return "该pdf没有内容"

    full_text_str="\n".join(full_text)

    print(f"pdf共{len(full_text_str)}字,正在存入数据库")
    result = kb_service.upload_by_str(full_text_str,filename)

    return result


def word_service(file_path: str, filename: str, operator: str = "admin"):
    print(f"开始解析 Word: {filename}, 上传者: {operator}")
    full_text=[]

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        return f"word文档打开失败"

    for para_num,paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()#去除首尾空格和换行

        if text:
            full_text.append(text)
            print(f"已读取第 {para_num} 段")

    for table_num,table in enumerate(doc.tables, start=1):

        for row in table.rows:
            row_text =[]

            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell.text:
                    row_text.append(cell_text)

            if row_text:
                full_text.append("|".join(row_text))

    if not full_text:
        return f"word里没有文字"

    full_text_str = "\n".join(full_text)
    print(f"Word文档共 {len(full_text_str)} 字，开始存入知识库...")
    result = kb_service.upload_by_str(full_text_str,filename)

    return result



